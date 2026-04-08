"""
Document Loader Module
Load and extract text from various document types
"""

import asyncio
from pathlib import Path
from typing import Dict, Any, Optional
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd
from PIL import Image
import io
import aiofiles

from app.document_processor.ocr import OCRProcessor
from app.utils.logger import logger

class DocumentLoader:
    """
    Load and extract text from different document formats
    """
    
    def __init__(self):
        self.ocr = OCRProcessor()
        self.supported_extensions = {
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.doc': self._load_doc,
            '.xlsx': self._load_excel,
            '.xls': self._load_excel,
            '.txt': self._load_text,
            '.jpg': self._load_image,
            '.jpeg': self._load_image,
            '.png': self._load_image,
            '.gif': self._load_image
        }
    
    async def load(self, file_path: Path) -> Dict[str, Any]:
        """
        Load document from file path
        """
        logger.info(f"Loading document: {file_path}")
        
        # Get extension
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {extension}")
        
        # Load using appropriate method
        loader = self.supported_extensions[extension]
        result = await loader(file_path)
        
        logger.info(f"Document loaded: {result['metadata'].get('pages', 1)} pages, {len(result['text'])} chars")
        
        return result
    
    async def load_bytes(self, content: bytes, filename: str) -> Dict[str, Any]:
        """
        Load document from bytes
        """
        import tempfile
        import os
        
        # Create temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp:
            tmp.write(content)
            tmp_path = Path(tmp.name)
        
        try:
            return await self.load(tmp_path)
        finally:
            # Clean up
            if tmp_path.exists():
                tmp_path.unlink()
    
    async def _load_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from PDF
        """
        text_content = []
        pages = []
        
        try:
            # Try pdfplumber first
            with pdfplumber.open(file_path) as pdf:
                metadata = {
                    'pages': len(pdf.pages),
                    'metadata': pdf.metadata or {}
                }
                
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    text_content.append(text)
                    pages.append({
                        'number': i + 1,
                        'text': text,
                        'words': len(text.split())
                    })
                    
        except Exception as e:
            logger.warning(f"pdfplumber failed, falling back to PyPDF2: {e}")
            
            # Fallback to PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                metadata = {
                    'pages': len(reader.pages),
                    'metadata': reader.metadata
                }
                
                for i, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    text_content.append(text)
                    pages.append({
                        'number': i + 1,
                        'text': text,
                        'words': len(text.split())
                    })
        
        full_text = "\n".join(text_content)
        
        return {
            'text': full_text,
            'metadata': metadata,
            'pages': pages,
            'type': 'pdf'
        }
    
    async def _load_docx(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from DOCX
        """
        doc = DocxDocument(file_path)
        
        text_content = []
        for para in doc.paragraphs:
            if para.text:
                text_content.append(para.text)
        
        full_text = "\n".join(text_content)
        
        return {
            'text': full_text,
            'metadata': {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            },
            'type': 'docx'
        }
    
    async def _load_doc(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from older DOC format
        """
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                text = content.decode('utf-8', errors='ignore')
            
            return {
                'text': text,
                'metadata': {},
                'type': 'doc'
            }
        except Exception as e:
            logger.error(f"Failed to read DOC file: {e}")
            return {
                'text': "",
                'metadata': {'error': str(e)},
                'type': 'doc'
            }
    
    async def _load_excel(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from Excel
        """
        df_dict = pd.read_excel(file_path, sheet_name=None)
        
        text_content = []
        sheets = []
        
        for sheet_name, df in df_dict.items():
            sheet_text = f"Sheet: {sheet_name}\n"
            sheet_text += df.to_string()
            text_content.append(sheet_text)
            
            sheets.append({
                'name': sheet_name,
                'rows': len(df),
                'columns': len(df.columns)
            })
        
        return {
            'text': "\n\n".join(text_content),
            'metadata': {
                'sheets': sheets,
                'total_sheets': len(sheets)
            },
            'type': 'excel'
        }
    
    async def _load_text(self, file_path: Path) -> Dict[str, Any]:
        """
        Load text file
        """
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            text = await f.read()
        
        return {
            'text': text,
            'metadata': {
                'size': len(text),
                'lines': len(text.split('\n'))
            },
            'type': 'text'
        }
    
    async def _load_image(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from image using OCR
        """
        image = Image.open(file_path)
        
        text = await self.ocr.process_image(image)
        
        return {
            'text': text,
            'metadata': {
                'size': image.size,
                'format': image.format,
                'mode': image.mode
            },
            'type': 'image',
            'ocr_used': True
        }
    
    async def get_preview(self, file_path: Path, max_length: int = 1000) -> str:
        """
        Get document preview
        """
        result = await self.load(file_path)
        return result['text'][:max_length] + "..." if len(result['text']) > max_length else result['text']
