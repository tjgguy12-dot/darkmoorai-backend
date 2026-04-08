"""
Professional Invoice Builder with Logo Upload
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image
import requests
from io import BytesIO
import base64
import uuid
from datetime import datetime, timedelta
import qrcode
from pathlib import Path

class InvoiceBuilder:
    """Professional invoice builder with logo and multiple templates"""
    
    def __init__(self, temp_dir):
        self.temp_dir = temp_dir
    
    def create_invoice(self, data: dict, template: str = "professional") -> str:
        """
        Create professional invoice
        Templates: professional, modern, minimal, corporate, simple
        """
        templates = {
            "professional": self._create_professional_invoice,
            "modern": self._create_modern_invoice,
            "minimal": self._create_minimal_invoice,
            "corporate": self._create_corporate_invoice,
            "simple": self._create_simple_invoice
        }
        
        creator = templates.get(template, self._create_professional_invoice)
        return creator(data)
    
    def _add_logo(self, doc, logo_data: dict, position: str = "left"):
        """Add company logo from URL or base64"""
        try:
            if logo_data.get("url"):
                response = requests.get(logo_data["url"], timeout=10)
                img = Image.open(BytesIO(response.content))
            elif logo_data.get("base64"):
                img_data = base64.b64decode(logo_data["base64"])
                img = Image.open(BytesIO(img_data))
            elif logo_data.get("path"):
                img = Image.open(logo_data["path"])
            else:
                return None
            
            # Resize to professional dimensions
            img.thumbnail((120, 120))
            
            # Save temporarily
            temp_img = self.temp_dir / f"logo_{uuid.uuid4().hex[:8]}.png"
            img.save(temp_img)
            
            # Add to document
            paragraph = doc.add_paragraph()
            if position == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(temp_img), width=Inches(1.2))
            return paragraph
        except Exception as e:
            print(f"Logo addition failed: {e}")
            return None
    
    def _add_qr_code(self, doc, payment_data: dict):
        """Add QR code for payment"""
        try:
            # Create payment QR data
            qr_data = f"Pay to: {payment_data.get('bank_name', '')}\n"
            qr_data += f"Account: {payment_data.get('account_number', '')}\n"
            qr_data += f"Amount: {payment_data.get('total', '')}\n"
            qr_data += f"Reference: INV-{payment_data.get('invoice_number', '')}"
            
            qr = qrcode.make(qr_data)
            temp_qr = self.temp_dir / f"qr_{uuid.uuid4().hex[:8]}.png"
            qr.save(temp_qr)
            
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(temp_qr), width=Inches(1.5))
            doc.add_paragraph("Scan to pay")
            return paragraph
        except Exception as e:
            print(f"QR generation failed: {e}")
            return None
    
    def _create_professional_invoice(self, data: dict) -> str:
        """Professional invoice template with company branding"""
        doc = Document()
        
        # Add logo
        if data.get("company_logo"):
            self._add_logo(doc, data["company_logo"], "left")
        
        # Company header
        company = doc.add_paragraph()
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if data.get("company_name"):
            company.add_run(data["company_name"]).bold = True
            company.add_run("\n")
        if data.get("company_address"):
            company.add_run(data["company_address"])
            company.add_run("\n")
        if data.get("company_email") or data.get("company_phone"):
            contact = []
            if data.get("company_email"):
                contact.append(data["company_email"])
            if data.get("company_phone"):
                contact.append(data["company_phone"])
            if data.get("company_website"):
                contact.append(data["company_website"])
            company.add_run(" | ".join(contact))
        
        doc.add_paragraph()
        
        # Invoice title
        title = doc.add_heading("INVOICE", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0, 112, 192)
        
        # Invoice details table
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # Left side - Bill To
        bill_to = table.cell(0, 0)
        bill_to.paragraphs[0].add_run("BILL TO:").bold = True
        if data.get("client_name"):
            bill_to.add_paragraph(data["client_name"])
        if data.get("client_email"):
            bill_to.add_paragraph(data["client_email"])
        if data.get("client_address"):
            bill_to.add_paragraph(data["client_address"])
        if data.get("client_phone"):
            bill_to.add_paragraph(data["client_phone"])
        
        # Right side - Invoice Info
        invoice_info = table.cell(0, 1)
        invoice_info.paragraphs[0].add_run("INVOICE INFO:").bold = True
        
        invoice_number = data.get("invoice_number", f"INV-{datetime.now().strftime('%Y%m%d')}-001")
        invoice_info.add_paragraph(f"Invoice #: {invoice_number}")
        
        invoice_date = data.get("invoice_date", datetime.now().strftime("%Y-%m-%d"))
        invoice_info.add_paragraph(f"Date: {invoice_date}")
        
        due_date = data.get("due_date", (datetime.now() + timedelta(days=30)).strftime("%Y-%m-%d"))
        invoice_info.add_paragraph(f"Due Date: {due_date}")
        
        if data.get("po_number"):
            invoice_info.add_paragraph(f"PO Number: {data['po_number']}")
        
        # Payment terms
        payment_terms = table.cell(1, 0)
        payment_terms.paragraphs[0].add_run("PAYMENT TERMS:").bold = True
        terms = data.get("payment_terms", "Net 30 days")
        payment_terms.add_paragraph(terms)
        
        # Bank details
        bank_info = table.cell(1, 1)
        bank_info.paragraphs[0].add_run("BANK DETAILS:").bold = True
        if data.get("bank_name"):
            bank_info.add_paragraph(f"Bank: {data['bank_name']}")
        if data.get("account_name"):
            bank_info.add_paragraph(f"Account Name: {data['account_name']}")
        if data.get("account_number"):
            bank_info.add_paragraph(f"Account Number: {data['account_number']}")
        if data.get("routing_number"):
            bank_info.add_paragraph(f"Routing #: {data['routing_number']}")
        
        doc.add_paragraph()
        
        # Items table
        doc.add_heading("Items", level=2)
        items_table = doc.add_table(rows=1, cols=5)
        items_table.style = 'Table Grid'
        
        # Headers
        headers = ["#", "Description", "Quantity", "Unit Price", "Total"]
        for i, header in enumerate(headers):
            cell = items_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        # Add items
        subtotal = 0
        for idx, item in enumerate(data.get("items", []), 1):
            row = items_table.add_row().cells
            row[0].text = str(idx)
            row[1].text = item.get("description", "")
            row[2].text = str(item.get("quantity", 1))
            row[3].text = f"${item.get('unit_price', 0):,.2f}"
            item_total = item.get("quantity", 1) * item.get("unit_price", 0)
            row[4].text = f"${item_total:,.2f}"
            subtotal += item_total
        
        # Totals section
        doc.add_paragraph()
        totals_table = doc.add_table(rows=5, cols=2)
        totals_table.autofit = False
        totals_table.columns[0].width = Inches(5)
        totals_table.columns[1].width = Inches(2)
        
        # Subtotal
        totals_table.cell(0, 0).text = "Subtotal:"
        totals_table.cell(0, 1).text = f"${subtotal:,.2f}"
        
        # Tax
        tax_rate = data.get("tax_rate", 0)
        tax_amount = subtotal * (tax_rate / 100)
        totals_table.cell(1, 0).text = f"Tax ({tax_rate}%):"
        totals_table.cell(1, 1).text = f"${tax_amount:,.2f}"
        
        # Discount
        discount = data.get("discount", 0)
        totals_table.cell(2, 0).text = f"Discount:"
        totals_table.cell(2, 1).text = f"-${discount:,.2f}"
        
        # Shipping
        shipping = data.get("shipping", 0)
        totals_table.cell(3, 0).text = "Shipping:"
        totals_table.cell(3, 1).text = f"${shipping:,.2f}"
        
        # Total
        total = subtotal + tax_amount - discount + shipping
        totals_table.cell(4, 0).text = "TOTAL:"
        totals_table.cell(4, 1).text = f"${total:,.2f}"
        totals_table.cell(4, 0).paragraphs[0].runs[0].bold = True
        totals_table.cell(4, 1).paragraphs[0].runs[0].bold = True
        
        # Add QR code for payment
        if data.get("enable_qr"):
            payment_data = {
                "bank_name": data.get("bank_name", ""),
                "account_number": data.get("account_number", ""),
                "total": total,
                "invoice_number": invoice_number
            }
            self._add_qr_code(doc, payment_data)
        
        # Notes
        if data.get("notes"):
            doc.add_paragraph()
            doc.add_heading("Notes", level=2)
            doc.add_paragraph(data["notes"])
        
        # Footer
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Thank you for your business!").italic = True
        
        filename = f"invoice_{invoice_number}.docx"
        filepath = self.temp_dir / filename
        doc.save(filepath)
        return str(filepath)
    
    def _create_modern_invoice(self, data: dict) -> str:
        """Modern colorful invoice template"""
        doc = Document()
        
        # Color bar
        doc.add_paragraph("_" * 80)
        
        # Logo and company
        if data.get("company_logo"):
            self._add_logo(doc, data["company_logo"], "left")
        
        # Company name
        if data.get("company_name"):
            name = doc.add_heading(data["company_name"], level=1)
            for run in name.runs:
                run.font.color.rgb = RGBColor(41, 128, 185)
        
        # Invoice title with color
        title = doc.add_heading("INVOICE", level=2)
        for run in title.runs:
            run.font.color.rgb = RGBColor(231, 76, 60)
        
        # Two-column layout
        table = doc.add_table(rows=2, cols=2)
        
        # Client info
        table.cell(0, 0).paragraphs[0].add_run("Client:").bold = True
        if data.get("client_name"):
            table.cell(0, 0).add_paragraph(data["client_name"])
        if data.get("client_email"):
            table.cell(0, 0).add_paragraph(data["client_email"])
        
        # Invoice details
        table.cell(0, 1).paragraphs[0].add_run("Invoice Details:").bold = True
        invoice_number = data.get("invoice_number", f"INV-{datetime.now().strftime('%Y%m%d')}-001")
        table.cell(0, 1).add_paragraph(f"#: {invoice_number}")
        table.cell(0, 1).add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        
        # Items table with modern styling
        items_table = doc.add_table(rows=1, cols=4)
        items_table.style = 'Table Grid'
        
        headers = ["Item", "Qty", "Price", "Total"]
        for i, header in enumerate(headers):
            cell = items_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            # Color header background
            shading = cell._tc.get_or_add_tcPr()
            shading.append(OxmlElement('w:shd'))
        
        # Items
        for item in data.get("items", []):
            row = items_table.add_row().cells
            row[0].text = item.get("description", "")
            row[1].text = str(item.get("quantity", 1))
            row[2].text = f"${item.get('unit_price', 0):,.2f}"
            total = item.get("quantity", 1) * item.get("unit_price", 0)
            row[3].text = f"${total:,.2f}"
        
        filename = f"invoice_modern_{invoice_number}.docx"
        filepath = self.temp_dir / filename
        doc.save(filepath)
        return str(filepath)
    
    def _create_minimal_invoice(self, data: dict) -> str:
        """Clean minimal invoice template"""
        doc = Document()
        
        # Company
        if data.get("company_name"):
            doc.add_heading(data["company_name"], level=0)
        else:
            doc.add_heading("INVOICE", level=0)
        
        doc.add_paragraph()
        
        # Client and invoice info side by side
        table = doc.add_table(rows=1, cols=2)
        
        # Left
        table.cell(0, 0).add_paragraph(f"Bill to:\n{data.get('client_name', '')}\n{data.get('client_email', '')}")
        
        # Right
        invoice_number = data.get("invoice_number", f"INV-{datetime.now().strftime('%Y%m%d')}-001")
        table.cell(0, 1).add_paragraph(f"Invoice #: {invoice_number}\nDate: {datetime.now().strftime('%Y-%m-%d')}")
        
        # Simple items table
        items_table = doc.add_table(rows=1, cols=3)
        headers = ["Description", "Qty", "Amount"]
        for i, header in enumerate(headers):
            items_table.rows[0].cells[i].text = header
        
        for item in data.get("items", []):
            row = items_table.add_row().cells
            row[0].text = item.get("description", "")
            row[1].text = str(item.get("quantity", 1))
            total = item.get("quantity", 1) * item.get("unit_price", 0)
            row[2].text = f"${total:,.2f}"
        
        # Total
        total = sum(item.get("quantity", 1) * item.get("unit_price", 0) for item in data.get("items", []))
        doc.add_paragraph(f"Total: ${total:,.2f}")
        
        filename = f"invoice_minimal_{invoice_number}.docx"
        filepath = self.temp_dir / filename
        doc.save(filepath)
        return str(filepath)
    
    def _create_corporate_invoice(self, data: dict) -> str:
        """Corporate invoice with professional layout"""
        doc = Document()
        
        # Header with border
        header_table = doc.add_table(rows=1, cols=2)
        header_table.style = 'Table Grid'
        
        # Left side - Company
        left_cell = header_table.cell(0, 0)
        if data.get("company_logo"):
            self._add_logo(doc, data["company_logo"], "left")
        if data.get("company_name"):
            left_cell.add_paragraph(data["company_name"]).bold = True
        if data.get("company_address"):
            left_cell.add_paragraph(data["company_address"])
        
        # Right side - Invoice
        right_cell = header_table.cell(0, 1)
        right_cell.paragraphs[0].add_run("INVOICE").bold = True
        invoice_number = data.get("invoice_number", f"INV-{datetime.now().strftime('%Y%m%d')}-001")
        right_cell.add_paragraph(f"Number: {invoice_number}")
        right_cell.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        
        doc.add_paragraph()
        
        # Client section
        doc.add_heading("CLIENT INFORMATION", level=2)
        doc.add_paragraph(f"Name: {data.get('client_name', '')}")
        doc.add_paragraph(f"Email: {data.get('client_email', '')}")
        if data.get("client_phone"):
            doc.add_paragraph(f"Phone: {data['client_phone']}")
        
        # Items
        doc.add_heading("INVOICE DETAILS", level=2)
        items_table = doc.add_table(rows=1, cols=5)
        items_table.style = 'Table Grid'
        
        headers = ["#", "Item Description", "Quantity", "Unit Price", "Line Total"]
        for i, header in enumerate(headers):
            items_table.rows[0].cells[i].text = header
        
        for idx, item in enumerate(data.get("items", []), 1):
            row = items_table.add_row().cells
            row[0].text = str(idx)
            row[1].text = item.get("description", "")
            row[2].text = str(item.get("quantity", 1))
            row[3].text = f"${item.get('unit_price', 0):,.2f}"
            total = item.get("quantity", 1) * item.get("unit_price", 0)
            row[4].text = f"${total:,.2f}"
        
        filename = f"invoice_corporate_{invoice_number}.docx"
        filepath = self.temp_dir / filename
        doc.save(filepath)
        return str(filepath)
    
    def _create_simple_invoice(self, data: dict) -> str:
        """Simple invoice template"""
        doc = Document()
        
        doc.add_heading("INVOICE", level=0)
        doc.add_paragraph(f"Invoice #: {data.get('invoice_number', f'INV-{datetime.now().strftime("%Y%m%d")}-001')}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        doc.add_paragraph()
        doc.add_paragraph(f"Client: {data.get('client_name', '')}")
        doc.add_paragraph()
        
        for item in data.get("items", []):
            doc.add_paragraph(f"{item.get('description', '')} - {item.get('quantity', 1)} x ${item.get('unit_price', 0):,.2f} = ${item.get('quantity', 1) * item.get('unit_price', 0):,.2f}")
        
        total = sum(item.get("quantity", 1) * item.get("unit_price", 0) for item in data.get("items", []))
        doc.add_paragraph()
        doc.add_paragraph(f"TOTAL: ${total:,.2f}").bold = True
        
        filename = f"invoice_simple_{data.get('invoice_number', '001')}.docx"
        filepath = self.temp_dir / filename
        doc.save(filepath)
        return str(filepath)


invoice_builder = InvoiceBuilder