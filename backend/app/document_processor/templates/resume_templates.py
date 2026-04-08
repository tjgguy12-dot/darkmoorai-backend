"""
Professional Resume Templates with Photo Support
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import requests
from io import BytesIO
import base64
from pathlib import Path
import uuid

class ResumeBuilder:
    """Professional resume builder with photo and multiple templates"""
    
    def __init__(self, temp_dir):
        self.temp_dir = temp_dir
    
    def create_resume(self, data: dict, template: str = "modern") -> str:
        """
        Create professional resume
        Templates: modern, classic, creative, executive, tech
        """
        templates = {
            "modern": self._create_modern_resume,
            "classic": self._create_classic_resume,
            "creative": self._create_creative_resume,
            "executive": self._create_executive_resume,
            "tech": self._create_tech_resume
        }
        
        creator = templates.get(template, self._create_modern_resume)
        return creator(data)
    
    def _add_photo(self, doc, photo_data: dict, position: str = "right"):
        """Add photo to document from URL or base64"""
        try:
            if photo_data.get("url"):
                response = requests.get(photo_data["url"], timeout=10)
                img = Image.open(BytesIO(response.content))
            elif photo_data.get("base64"):
                img_data = base64.b64decode(photo_data["base64"])
                img = Image.open(BytesIO(img_data))
            elif photo_data.get("path"):
                img = Image.open(photo_data["path"])
            else:
                return None
            
            # Resize to professional dimensions (2x2 inches)
            img.thumbnail((150, 150))
            
            # Save temporarily
            temp_img = self.temp_dir / f"photo_{uuid.uuid4().hex[:8]}.png"
            img.save(temp_img)
            
            # Add to document
            if position == "right":
                # Create table for header with photo
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Inches(4.5)
                table.columns[1].width = Inches(1.5)
                
                # Name in left cell
                name_cell = table.cell(0, 0)
                # Photo in right cell
                photo_cell = table.cell(0, 1)
                photo_cell.paragraphs[0].add_run().add_picture(str(temp_img), width=Inches(1.2))
                photo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                return table
            else:
                # Center photo at top
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(temp_img), width=Inches(1.5))
                return None
                
        except Exception as e:
            print(f"Photo addition failed: {e}")
            return None
    
    def _create_modern_resume(self, data: dict) -> str:
        """Modern template with sidebar"""
        doc = Document()
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        
        # Add photo if provided
        if data.get("photo"):
            self._add_photo(doc, data["photo"], "right")
        
        # Name
        name = doc.add_heading(data.get("name", "Your Name"), 0)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title
        title = doc.add_paragraph(data.get("title", "Professional"))
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].italic = True
        
        # Contact info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info = []
        if data.get("email"):
            contact_info.append(f"✉ {data['email']}")
        if data.get("phone"):
            contact_info.append(f"📞 {data['phone']}")
        if data.get("location"):
            contact_info.append(f"📍 {data['location']}")
        if data.get("linkedin"):
            contact_info.append(f"🔗 {data['linkedin']}")
        if data.get("github"):
            contact_info.append(f"💻 {data['github']}")
        contact.add_run(" | ".join(contact_info))
        doc.add_paragraph()
        
        # Two-column layout
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(4.5)
        
        # LEFT COLUMN
        left_cell = table.cell(0, 0)
        
        # Skills
        if data.get("skills"):
            heading = left_cell.add_heading("SKILLS", level=2)
            self._style_heading(heading, "modern")
            for skill in data["skills"]:
                p = left_cell.add_paragraph()
                p.add_run("▹ ").bold = True
                if isinstance(skill, dict):
                    p.add_run(f"{skill.get('name', '')}")
                    if skill.get("level"):
                        p.add_run(f" - {skill.get('level')}")
                else:
                    p.add_run(str(skill))
            left_cell.add_paragraph()
        
        # Languages
        if data.get("languages"):
            heading = left_cell.add_heading("LANGUAGES", level=2)
            self._style_heading(heading, "modern")
            for lang in data["languages"]:
                p = left_cell.add_paragraph()
                p.add_run("▹ ").bold = True
                if isinstance(lang, dict):
                    p.add_run(f"{lang.get('language', '')} - {lang.get('level', '')}")
                else:
                    p.add_run(str(lang))
            left_cell.add_paragraph()
        
        # Certifications
        if data.get("certifications"):
            heading = left_cell.add_heading("CERTIFICATIONS", level=2)
            self._style_heading(heading, "modern")
            for cert in data["certifications"]:
                p = left_cell.add_paragraph()
                p.add_run("▹ ").bold = True
                if isinstance(cert, dict):
                    p.add_run(f"{cert.get('name', '')}")
                    if cert.get("date"):
                        p.add_run(f" ({cert.get('date')})")
                else:
                    p.add_run(str(cert))
            left_cell.add_paragraph()
        
        # RIGHT COLUMN
        right_cell = table.cell(0, 1)
        
        # Professional Summary
        if data.get("summary"):
            heading = right_cell.add_heading("PROFESSIONAL SUMMARY", level=2)
            self._style_heading(heading, "modern")
            right_cell.add_paragraph(data["summary"])
            right_cell.add_paragraph()
        
        # Work Experience
        if data.get("experience"):
            heading = right_cell.add_heading("WORK EXPERIENCE", level=2)
            self._style_heading(heading, "modern")
            for exp in data["experience"]:
                p = right_cell.add_paragraph()
                p.add_run(exp.get("title", "")).bold = True
                p.add_run(f" - {exp.get('company', '')}")
                p = right_cell.add_paragraph()
                p.add_run(exp.get("date", "")).italic = True
                right_cell.add_paragraph(exp.get("description", ""))
                if exp.get("achievements"):
                    for achievement in exp["achievements"]:
                        right_cell.add_paragraph(f"• {achievement}", style='List Bullet')
            right_cell.add_paragraph()
        
        # Education
        if data.get("education"):
            heading = right_cell.add_heading("EDUCATION", level=2)
            self._style_heading(heading, "modern")
            for edu in data["education"]:
                p = right_cell.add_paragraph()
                p.add_run(edu.get("degree", "")).bold = True
                p.add_run(f" - {edu.get('school', '')}")
                if edu.get("date"):
                    right_cell.add_paragraph(edu.get("date")).italic = True
                if edu.get("gpa"):
                    right_cell.add_paragraph(f"GPA: {edu.get('gpa')}")
                if edu.get("courses"):
                    right_cell.add_paragraph(f"Relevant Courses: {', '.join(edu['courses'])}")
            right_cell.add_paragraph()
        
        # Projects
        if data.get("projects"):
            heading = right_cell.add_heading("PROJECTS", level=2)
            self._style_heading(heading, "modern")
            for project in data["projects"]:
                p = right_cell.add_paragraph()
                p.add_run(project.get("name", "")).bold = True
                if project.get("technologies"):
                    p.add_run(f" - {project.get('technologies')}")
                right_cell.add_paragraph(project.get("description", ""))
                if project.get("link"):
                    right_cell.add_paragraph(f"Link: {project.get('link')}")
        
        # Save document
        filename = f"resume_{data.get('name', 'resume').replace(' ', '_')}.docx"
        filepath = self.temp_dir / filename
        doc.save(filepath)
        return str(filepath)
    
    def _create_classic_resume(self, data: dict) -> str:
        """Classic traditional resume template"""
        doc = Document()
        
        # Header
        name = doc.add_heading(data.get("name", "Your Name"), 0)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact line
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info = []
        if data.get("email"):
            contact_info.append(data['email'])
        if data.get("phone"):
            contact_info.append(data['phone'])
        if data.get("location"):
            contact_info.append(data['location'])
        contact.add_run(" | ".join(contact_info))
        
        # Horizontal line
        doc.add_paragraph("_" * 60)
        
        # Summary
        if data.get("summary"):
            doc.add_heading("SUMMARY", level=1)
            doc.add_paragraph(data["summary"])
        
        # Experience
        if data.get("experience"):
            doc.add_heading("EXPERIENCE", level=1)
            for exp in data["experience"]:
                p = doc.add_paragraph()
                p.add_run(exp.get("title", "")).bold = True
                doc.add_paragraph(f"{exp.get('company', '')} | {exp.get('date', '')}")
                doc.add_paragraph(exp.get("description", ""))
        
        # Education
        if data.get("education"):
            doc.add_heading("EDUCATION", level=1)
            for edu in data["education"]:
                p = doc.add_paragraph()
                p.add_run(edu.get("degree", "")).bold = True
                doc.add_paragraph(edu.get("school", ""))
                if edu.get("date"):
                    doc.add_paragraph(edu.get("date"))
        
        # Skills
        if data.get("skills"):
            doc.add_heading("SKILLS", level=1)
            skills_text = ", ".join([s if isinstance(s, str) else s.get('name', '') for s in data["skills"]])
            doc.add_paragraph(skills_text)
        
        filename = f"resume_classic_{data.get('name', 'resume').replace(' ', '_')}.docx"
        filepath = self.temp_dir / filename
        doc.save(filepath)
        return str(filepath)
    
    def _create_creative_resume(self, data: dict) -> str:
        """Creative/designer resume with colors and icons"""
        doc = Document()
        
        # Color accent line
        doc.add_paragraph("_" * 80)
        
        # Name with color
        name = doc.add_heading(data.get("name", "Your Name"), 0)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in name.runs:
            run.font.color.rgb = RGBColor(41, 128, 185)
        
        # Creative title
        title = doc.add_paragraph(data.get("title", "Creative Professional"))
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(52, 73, 94)
        
        # Portfolio link if provided
        if data.get("portfolio"):
            doc.add_paragraph(f"Portfolio: {data['portfolio']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Skills with rating bars
        if data.get("skills"):
            doc.add_heading("Core Competencies", level=1)
            for skill in data["skills"]:
                skill_name = skill if isinstance(skill, str) else skill.get('name', '')
                skill_level = 70  # default
                if isinstance(skill, dict) and skill.get("level"):
                    level_map = {"Beginner": 30, "Intermediate": 60, "Advanced": 85, "Expert": 100}
                    skill_level = level_map.get(skill["level"], 70)
                
                p = doc.add_paragraph()
                p.add_run(f"{skill_name}: ").bold = True
                p.add_run("█" * (skill_level // 10) + "░" * (10 - (skill_level // 10)))
        
        doc.add_paragraph()
        
        # Experience with achievements
        if data.get("experience"):
            doc.add_heading("Selected Achievements", level=1)
            for exp in data["experience"]:
                doc.add_heading(exp.get("title", ""), level=2)
                doc.add_paragraph(f"{exp.get('company', '')} | {exp.get('date', '')}")
                if exp.get("achievements"):
                    for achievement in exp["achievements"]:
                        doc.add_paragraph(f"✨ {achievement}", style='List Bullet')
        
        filename = f"resume_creative_{data.get('name', 'resume').replace(' ', '_')}.docx"
        filepath = self.temp_dir / filename
        doc.save(filepath)
        return str(filepath)
    
    def _create_executive_resume(self, data: dict) -> str:
        """Executive-level resume with leadership focus"""
        doc = Document()
        
        # Header
        name = doc.add_heading(data.get("name", "Your Name"), 0)
        
        # Executive title
        title = doc.add_paragraph(data.get("title", "Executive Leader"))
        title.runs[0].bold = True
        
        # Contact
        contact = doc.add_paragraph()
        contact_info = []
        if data.get("email"):
            contact_info.append(data['email'])
        if data.get("phone"):
            contact_info.append(data['phone'])
        if data.get("linkedin"):
            contact_info.append(data['linkedin'])
        contact.add_run(" | ".join(contact_info))
        
        doc.add_paragraph()
        
        # Executive Summary
        if data.get("summary"):
            doc.add_heading("EXECUTIVE SUMMARY", level=1)
            doc.add_paragraph(data["summary"])
        
        # Leadership Experience
        if data.get("experience"):
            doc.add_heading("LEADERSHIP EXPERIENCE", level=1)
            for exp in data["experience"]:
                p = doc.add_paragraph()
                p.add_run(exp.get("title", "")).bold = True
                p.add_run(f" | {exp.get('company', '')}")
                doc.add_paragraph(exp.get("date", "")).italic = True
                doc.add_paragraph(exp.get("description", ""))
                
                # Key achievements as metrics
                if exp.get("metrics"):
                    for metric in exp["metrics"]:
                        doc.add_paragraph(f"✓ {metric}", style='List Bullet')
        
        # Core Competencies
        if data.get("skills"):
            doc.add_heading("CORE COMPETENCIES", level=1)
            skills_cols = []
            for skill in data["skills"]:
                skill_name = skill if isinstance(skill, str) else skill.get('name', '')
                skills_cols.append(skill_name)
            
            # Format as columns
            for i in range(0, len(skills_cols), 3):
                doc.add_paragraph(" | ".join(skills_cols[i:i+3]))
        
        filename = f"resume_executive_{data.get('name', 'resume').replace(' ', '_')}.docx"
        filepath = self.temp_dir / filename
        doc.save(filepath)
        return str(filepath)
    
    def _create_tech_resume(self, data: dict) -> str:
        """Tech/IT resume with technology focus"""
        doc = Document()
        
        # Header
        name = doc.add_heading(data.get("name", "Your Name"), 0)
        
        # Tech title
        title = doc.add_paragraph(data.get("title", "Software Engineer"))
        title.runs[0].bold = True
        
        # Contact with icons
        contact = doc.add_paragraph()
        contact_info = []
        if data.get("email"):
            contact_info.append(f"✉ {data['email']}")
        if data.get("phone"):
            contact_info.append(f"📱 {data['phone']}")
        if data.get("github"):
            contact_info.append(f"🐙 {data['github']}")
        if data.get("linkedin"):
            contact_info.append(f"🔗 {data['linkedin']}")
        contact.add_run("  |  ".join(contact_info))
        
        doc.add_paragraph()
        
        # Technical Skills with categories
        if data.get("skills"):
            doc.add_heading("TECHNICAL SKILLS", level=1)
            skill_categories = {"Languages": [], "Frameworks": [], "Tools": [], "Databases": []}
            
            for skill in data["skills"]:
                if isinstance(skill, dict):
                    category = skill.get("category", "Languages")
                    if category in skill_categories:
                        skill_categories[category].append(skill.get("name", ""))
                else:
                    skill_categories["Languages"].append(skill)
            
            for category, skills in skill_categories.items():
                if skills:
                    doc.add_heading(category, level=2)
                    doc.add_paragraph(", ".join(skills))
        
        # Experience with technologies
        if data.get("experience"):
            doc.add_heading("EXPERIENCE", level=1)
            for exp in data["experience"]:
                p = doc.add_paragraph()
                p.add_run(exp.get("title", "")).bold = True
                p.add_run(f" @ {exp.get('company', '')}")
                doc.add_paragraph(exp.get("date", "")).italic = True
                doc.add_paragraph(exp.get("description", ""))
                
                if exp.get("technologies"):
                    doc.add_paragraph(f"Tech Stack: {exp.get('technologies')}")
        
        # Projects with GitHub links
        if data.get("projects"):
            doc.add_heading("PROJECTS", level=1)
            for project in data["projects"]:
                p = doc.add_paragraph()
                p.add_run(project.get("name", "")).bold = True
                doc.add_paragraph(project.get("description", ""))
                if project.get("link"):
                    doc.add_paragraph(f"🔗 {project.get('link')}")
        
        filename = f"resume_tech_{data.get('name', 'resume').replace(' ', '_')}.docx"
        filepath = self.temp_dir / filename
        doc.save(filepath)
        return str(filepath)
    
    def _style_heading(self, heading, style: str):
        """Style heading based on template"""
        for run in heading.runs:
            run.font.bold = True
            if style == "modern":
                run.font.color.rgb = RGBColor(0, 112, 192)


resume_builder = ResumeBuilder